"""
File Processing System
Handles document upload, processing, and knowledge base integration
"""

import os
import io
import hashlib
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
import logging

# Document processing libraries
import PyPDF2
import docx
from PIL import Image
import pytesseract
import pandas as pd
import json

# Text processing
from sentence_transformers import SentenceTransformer

logger = logging.getLogger(__name__)

class FileProcessor:
    """Enhanced file processing system"""
    
    def __init__(self, upload_dir: str = "uploads", max_file_size: int = 50 * 1024 * 1024):  # 50MB
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        
        self.max_file_size = max_file_size
        self.supported_formats = {
            'text': ['.txt', '.md', '.rtf'],
            'pdf': ['.pdf'],
            'word': ['.doc', '.docx'],
            'excel': ['.xls', '.xlsx', '.csv'],
            'image': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff'],
            'audio': ['.wav', '.mp3', '.flac', '.m4a', '.ogg'],
            'json': ['.json'],
            'xml': ['.xml']
        }
        
        # Initialize text embeddings model
        self.embeddings_model = None
        self._initialize_embeddings()
    
    def _initialize_embeddings(self):
        """Initialize sentence transformer for text embeddings"""
        try:
            self.embeddings_model = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("✅ Text embeddings model loaded")
        except Exception as e:
            logger.error(f"❌ Failed to load embeddings model: {e}")
    
    async def process_uploaded_file(
        self, 
        file_content: bytes, 
        filename: str, 
        user_id: str,
        description: Optional[str] = None
    ) -> Dict[str, Any]:
        """Process uploaded file and extract content"""
        
        try:
            # Validate file
            validation_result = self._validate_file(file_content, filename)
            if not validation_result['valid']:
                return {
                    "success": False,
                    "error": validation_result['error'],
                    "file_info": None
                }
            
            # Generate file hash for deduplication
            file_hash = hashlib.md5(file_content).hexdigest()
            
            # Save file
            file_path = self._save_file(file_content, filename, user_id, file_hash)
            
            # Extract content based on file type
            extraction_result = await self._extract_content(file_path, filename)
            
            if not extraction_result['success']:
                return {
                    "success": False,
                    "error": extraction_result['error'],
                    "file_info": None
                }
            
            # Create file metadata
            file_info = {
                "file_id": file_hash,
                "original_filename": filename,
                "file_path": str(file_path),
                "user_id": user_id,
                "upload_timestamp": datetime.now().isoformat(),
                "file_size": len(file_content),
                "file_type": self._get_file_type(filename),
                "mime_type": mimetypes.guess_type(filename)[0],
                "description": description,
                "content": extraction_result['content'],
                "metadata": extraction_result['metadata'],
                "processing_time": extraction_result.get('processing_time', 0)
            }
            
            # Generate embeddings for text content
            if extraction_result['content']:
                embeddings = await self._generate_embeddings(extraction_result['content'])
                file_info['embeddings'] = embeddings
            
            # Save metadata
            await self._save_file_metadata(file_info)
            
            logger.info(f"✅ File processed successfully: {filename}")
            
            return {
                "success": True,
                "file_info": file_info,
                "content_preview": extraction_result['content'][:500] + "..." if len(extraction_result['content']) > 500 else extraction_result['content']
            }
            
        except Exception as e:
            logger.error(f"❌ File processing error: {e}")
            return {
                "success": False,
                "error": str(e),
                "file_info": None
            }
    
    def _validate_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Validate uploaded file"""
        
        # Check file size
        if len(file_content) > self.max_file_size:
            return {
                "valid": False,
                "error": f"File size ({len(file_content)} bytes) exceeds maximum allowed size ({self.max_file_size} bytes)"
            }
        
        # Check file extension
        file_ext = Path(filename).suffix.lower()
        supported_extensions = []
        for format_exts in self.supported_formats.values():
            supported_extensions.extend(format_exts)
        
        if file_ext not in supported_extensions:
            return {
                "valid": False,
                "error": f"File format '{file_ext}' not supported. Supported formats: {', '.join(supported_extensions)}"
            }
        
        # Check for empty file
        if len(file_content) == 0:
            return {
                "valid": False,
                "error": "File is empty"
            }
        
        return {"valid": True, "error": None}
    
    def _save_file(self, file_content: bytes, filename: str, user_id: str, file_hash: str) -> Path:
        """Save file to upload directory"""
        
        # Create user directory
        user_dir = self.upload_dir / user_id
        user_dir.mkdir(exist_ok=True)
        
        # Create filename with hash to avoid conflicts
        file_ext = Path(filename).suffix
        safe_filename = f"{file_hash}_{filename}"
        file_path = user_dir / safe_filename
        
        # Save file
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        return file_path
    
    async def _extract_content(self, file_path: Path, filename: str) -> Dict[str, Any]:
        """Extract content from file based on type"""
        
        start_time = datetime.now()
        file_type = self._get_file_type(filename)
        
        try:
            if file_type == 'text':
                content = await self._extract_text_content(file_path)
            elif file_type == 'pdf':
                content = await self._extract_pdf_content(file_path)
            elif file_type == 'word':
                content = await self._extract_word_content(file_path)
            elif file_type == 'excel':
                content = await self._extract_excel_content(file_path)
            elif file_type == 'image':
                content = await self._extract_image_content(file_path)
            elif file_type == 'json':
                content = await self._extract_json_content(file_path)
            else:
                return {
                    "success": False,
                    "error": f"Content extraction not implemented for file type: {file_type}",
                    "content": "",
                    "metadata": {}
                }
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return {
                "success": True,
                "content": content['text'],
                "metadata": content.get('metadata', {}),
                "processing_time": processing_time
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Content extraction failed: {e}",
                "content": "",
                "metadata": {}
            }
    
    async def _extract_text_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from text files"""
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        return {
            "text": content,
            "metadata": {
                "line_count": len(content.split('\n')),
                "word_count": len(content.split()),
                "character_count": len(content)
            }
        }
    
    async def _extract_pdf_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from PDF files"""
        
        text_content = ""
        metadata = {"page_count": 0}
        
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            metadata["page_count"] = len(pdf_reader.pages)
            
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
        
        metadata.update({
            "word_count": len(text_content.split()),
            "character_count": len(text_content)
        })
        
        return {"text": text_content, "metadata": metadata}
    
    async def _extract_word_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from Word documents"""
        
        doc = docx.Document(file_path)
        text_content = ""
        
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "word_count": len(text_content.split()),
            "character_count": len(text_content)
        }
        
        return {"text": text_content, "metadata": metadata}
    
    async def _extract_excel_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from Excel/CSV files"""
        
        if file_path.suffix.lower() == '.csv':
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)
        
        # Convert DataFrame to text representation
        text_content = df.to_string()
        
        metadata = {
            "row_count": len(df),
            "column_count": len(df.columns),
            "columns": list(df.columns),
            "data_types": df.dtypes.to_dict()
        }
        
        return {"text": text_content, "metadata": metadata}
    
    async def _extract_image_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from images using OCR"""
        
        try:
            # Use pytesseract for OCR
            image = Image.open(file_path)
            text_content = pytesseract.image_to_string(image)
            
            metadata = {
                "image_size": image.size,
                "image_mode": image.mode,
                "ocr_confidence": "estimated_high" if len(text_content.strip()) > 10 else "estimated_low"
            }
            
            return {"text": text_content, "metadata": metadata}
            
        except Exception as e:
            # If OCR fails, return image metadata only
            image = Image.open(file_path)
            return {
                "text": f"Image file: {file_path.name}",
                "metadata": {
                    "image_size": image.size,
                    "image_mode": image.mode,
                    "ocr_error": str(e)
                }
            }
    
    async def _extract_json_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from JSON files"""
        
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        # Convert JSON to readable text
        text_content = json.dumps(json_data, indent=2)
        
        metadata = {
            "json_structure": type(json_data).__name__,
            "key_count": len(json_data) if isinstance(json_data, dict) else None,
            "item_count": len(json_data) if isinstance(json_data, list) else None
        }
        
        return {"text": text_content, "metadata": metadata}
    
    async def _generate_embeddings(self, text: str) -> Optional[List[float]]:
        """Generate embeddings for text content"""
        
        if not self.embeddings_model or not text.strip():
            return None
        
        try:
            # Truncate text if too long (model limit)
            max_length = 512
            if len(text) > max_length:
                text = text[:max_length]
            
            embeddings = self.embeddings_model.encode(text)
            return embeddings.tolist()
            
        except Exception as e:
            logger.error(f"❌ Embeddings generation error: {e}")
            return None
    
    async def _save_file_metadata(self, file_info: Dict[str, Any]):
        """Save file metadata to JSON"""
        
        metadata_dir = self.upload_dir / "metadata"
        metadata_dir.mkdir(exist_ok=True)
        
        metadata_file = metadata_dir / f"{file_info['file_id']}.json"
        
        # Remove embeddings for JSON storage (too large)
        metadata_copy = file_info.copy()
        if 'embeddings' in metadata_copy:
            del metadata_copy['embeddings']
        
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata_copy, f, indent=2, ensure_ascii=False)
    
    def _get_file_type(self, filename: str) -> str:
        """Determine file type from extension"""
        
        file_ext = Path(filename).suffix.lower()
        
        for file_type, extensions in self.supported_formats.items():
            if file_ext in extensions:
                return file_type
        
        return 'unknown'
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get supported file formats"""
        return self.supported_formats.copy()
    
    def get_upload_stats(self) -> Dict[str, Any]:
        """Get file upload statistics"""
        
        total_files = 0
        total_size = 0
        file_types = {}
        
        for user_dir in self.upload_dir.iterdir():
            if user_dir.is_dir() and user_dir.name != "metadata":
                for file_path in user_dir.iterdir():
                    if file_path.is_file():
                        total_files += 1
                        total_size += file_path.stat().st_size
                        
                        file_type = self._get_file_type(file_path.name)
                        file_types[file_type] = file_types.get(file_type, 0) + 1
        
        return {
            "total_files": total_files,
            "total_size_bytes": total_size,
            "total_size_mb": round(total_size / (1024 * 1024), 2),
            "file_types": file_types,
            "upload_directory": str(self.upload_dir),
            "max_file_size_mb": round(self.max_file_size / (1024 * 1024), 2)
        }

# Global file processor instance
file_processor = FileProcessor()
