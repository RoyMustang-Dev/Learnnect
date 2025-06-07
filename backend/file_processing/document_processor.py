"""
Document Processing Module for Learnnect AI Chatbot
Handles PDF, Word, Text, CSV, JSON, and other file formats
"""

import asyncio
import json
import logging
import mimetypes
import re
from datetime import datetime
from typing import Dict, Any, List, Optional, Union
import pandas as pd
import PyPDF2
from docx import Document
import csv
import io

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        """Initialize document processor"""
        
        # Supported file types
        self.supported_types = {
            'text/plain': self.process_text,
            'application/pdf': self.process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self.process_docx,
            'application/msword': self.process_doc,
            'text/csv': self.process_csv,
            'application/json': self.process_json,
            'application/vnd.ms-excel': self.process_excel,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self.process_excel
        }
        
        logger.info("📄 Document Processor initialized")
    
    async def process_document(self, 
                             file_content: bytes, 
                             filename: str, 
                             content_type: str = None) -> Dict[str, Any]:
        """Main document processing method"""
        
        start_time = datetime.now()
        
        try:
            # Detect content type if not provided
            if not content_type:
                content_type, _ = mimetypes.guess_type(filename)
            
            logger.info(f"📄 Processing document: {filename} ({content_type})")
            
            # Check if file type is supported
            if content_type not in self.supported_types:
                raise ValueError(f"Unsupported file type: {content_type}")
            
            # Process document based on type
            processor = self.supported_types[content_type]
            extracted_data = await processor(file_content, filename)
            
            # Generate summary and key points
            summary = self.generate_summary(extracted_data['text'])
            key_points = self.extract_key_points(extracted_data['text'])
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            result = {
                'filename': filename,
                'document_type': content_type,
                'extracted_text': extracted_data['text'],
                'metadata': extracted_data.get('metadata', {}),
                'summary': summary,
                'key_points': key_points,
                'word_count': len(extracted_data['text'].split()),
                'character_count': len(extracted_data['text']),
                'processing_time': processing_time,
                'timestamp': datetime.now().isoformat()
            }
            
            logger.info(f"✅ Document processed: {len(extracted_data['text'])} characters")
            return result
            
        except Exception as e:
            logger.error(f"❌ Document processing error: {e}")
            raise
    
    async def process_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process plain text files"""
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Could not decode text file")
            
            return {
                'text': text.strip(),
                'metadata': {
                    'encoding': encoding,
                    'line_count': len(text.split('\n'))
                }
            }
            
        except Exception as e:
            logger.error(f"❌ Text processing error: {e}")
            raise
    
    async def process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF files"""
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            metadata = {
                'page_count': len(pdf_reader.pages),
                'title': '',
                'author': '',
                'subject': ''
            }
            
            # Extract metadata if available
            if pdf_reader.metadata:
                metadata.update({
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'subject': pdf_reader.metadata.get('/Subject', ''),
                    'creator': pdf_reader.metadata.get('/Creator', ''),
                    'producer': pdf_reader.metadata.get('/Producer', '')
                })
            
            # Extract text from all pages
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"⚠️ Could not extract text from page {page_num + 1}: {e}")
            
            full_text = '\n\n'.join(text_content)
            
            return {
                'text': full_text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"❌ PDF processing error: {e}")
            raise
    
    async def process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process Word DOCX files"""
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(' | '.join(row_data))
                tables_text.append('\n'.join(table_data))
            
            # Combine all text
            full_text = '\n\n'.join(paragraphs)
            if tables_text:
                full_text += '\n\n--- Tables ---\n\n' + '\n\n'.join(tables_text)
            
            metadata = {
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables)
            }
            
            return {
                'text': full_text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"❌ DOCX processing error: {e}")
            raise
    
    async def process_doc(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process legacy Word DOC files"""
        
        # For DOC files, we'll need python-docx2txt or similar
        # For now, return an error message
        raise ValueError("Legacy DOC files not supported. Please convert to DOCX format.")
    
    async def process_csv(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process CSV files"""
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    csv_text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Could not decode CSV file")
            
            # Parse CSV
            csv_file = io.StringIO(csv_text)
            csv_reader = csv.reader(csv_file)
            
            rows = list(csv_reader)
            if not rows:
                raise ValueError("Empty CSV file")
            
            headers = rows[0]
            data_rows = rows[1:]
            
            # Create readable text representation
            text_parts = [f"CSV Data from {filename}"]
            text_parts.append(f"Headers: {', '.join(headers)}")
            text_parts.append(f"Total rows: {len(data_rows)}")
            text_parts.append("")
            
            # Add sample data (first 10 rows)
            text_parts.append("Sample data:")
            for i, row in enumerate(data_rows[:10]):
                row_text = []
                for j, value in enumerate(row):
                    if j < len(headers):
                        row_text.append(f"{headers[j]}: {value}")
                text_parts.append(f"Row {i+1}: {', '.join(row_text)}")
            
            if len(data_rows) > 10:
                text_parts.append(f"... and {len(data_rows) - 10} more rows")
            
            # Add basic statistics
            text_parts.append("")
            text_parts.append("Column analysis:")
            for i, header in enumerate(headers):
                column_values = [row[i] if i < len(row) else '' for row in data_rows]
                non_empty = [v for v in column_values if v.strip()]
                text_parts.append(f"{header}: {len(non_empty)} non-empty values")
            
            full_text = '\n'.join(text_parts)
            
            metadata = {
                'row_count': len(data_rows),
                'column_count': len(headers),
                'headers': headers,
                'encoding': encoding
            }
            
            return {
                'text': full_text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"❌ CSV processing error: {e}")
            raise
    
    async def process_json(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process JSON files"""
        
        try:
            json_text = file_content.decode('utf-8')
            data = json.loads(json_text)
            
            # Create readable text representation
            text_parts = [f"JSON Data from {filename}"]
            text_parts.append("")
            
            # Analyze structure
            def analyze_json(obj, prefix="", max_depth=3, current_depth=0):
                if current_depth > max_depth:
                    return ["... (truncated)"]
                
                lines = []
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            lines.append(f"{prefix}{key}: {type(value).__name__}")
                            lines.extend(analyze_json(value, prefix + "  ", max_depth, current_depth + 1))
                        else:
                            lines.append(f"{prefix}{key}: {value}")
                elif isinstance(obj, list):
                    lines.append(f"{prefix}Array with {len(obj)} items")
                    if obj and current_depth < max_depth:
                        lines.append(f"{prefix}Sample item:")
                        lines.extend(analyze_json(obj[0], prefix + "  ", max_depth, current_depth + 1))
                else:
                    lines.append(f"{prefix}{obj}")
                
                return lines
            
            analysis = analyze_json(data)
            text_parts.extend(analysis)
            
            # Add raw JSON (truncated if too long)
            json_str = json.dumps(data, indent=2)
            if len(json_str) > 2000:
                text_parts.append("")
                text_parts.append("Raw JSON (truncated):")
                text_parts.append(json_str[:2000] + "...")
            else:
                text_parts.append("")
                text_parts.append("Raw JSON:")
                text_parts.append(json_str)
            
            full_text = '\n'.join(text_parts)
            
            metadata = {
                'json_type': type(data).__name__,
                'size_bytes': len(json_text),
                'structure_depth': self.get_json_depth(data)
            }
            
            return {
                'text': full_text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"❌ JSON processing error: {e}")
            raise
    
    async def process_excel(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process Excel files"""
        
        try:
            excel_file = io.BytesIO(file_content)
            
            # Read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            
            text_parts = [f"Excel Data from {filename}"]
            text_parts.append(f"Number of sheets: {len(excel_data)}")
            text_parts.append("")
            
            total_rows = 0
            total_cols = 0
            
            for sheet_name, df in excel_data.items():
                text_parts.append(f"Sheet: {sheet_name}")
                text_parts.append(f"Dimensions: {df.shape[0]} rows × {df.shape[1]} columns")
                text_parts.append(f"Columns: {', '.join(df.columns.astype(str))}")
                
                # Add sample data
                if not df.empty:
                    text_parts.append("Sample data:")
                    sample_rows = min(5, len(df))
                    for i in range(sample_rows):
                        row_data = []
                        for col in df.columns:
                            value = df.iloc[i][col]
                            row_data.append(f"{col}: {value}")
                        text_parts.append(f"  Row {i+1}: {', '.join(row_data)}")
                
                text_parts.append("")
                total_rows += df.shape[0]
                total_cols += df.shape[1]
            
            full_text = '\n'.join(text_parts)
            
            metadata = {
                'sheet_count': len(excel_data),
                'total_rows': total_rows,
                'total_columns': total_cols,
                'sheet_names': list(excel_data.keys())
            }
            
            return {
                'text': full_text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"❌ Excel processing error: {e}")
            raise
    
    def generate_summary(self, text: str, max_length: int = 200) -> str:
        """Generate a summary of the document"""
        
        if len(text) <= max_length:
            return text
        
        # Simple extractive summary - take first few sentences
        sentences = re.split(r'[.!?]+', text)
        summary_parts = []
        current_length = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and current_length + len(sentence) <= max_length:
                summary_parts.append(sentence)
                current_length += len(sentence)
            else:
                break
        
        summary = '. '.join(summary_parts)
        if len(summary) < len(text):
            summary += "..."
        
        return summary
    
    def extract_key_points(self, text: str, max_points: int = 5) -> List[str]:
        """Extract key points from the document"""
        
        # Simple key point extraction based on sentence length and position
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 20]
        
        # Score sentences based on length and keywords
        scored_sentences = []
        keywords = ['important', 'key', 'main', 'primary', 'essential', 'critical', 'significant']
        
        for sentence in sentences:
            score = len(sentence.split())  # Base score on word count
            
            # Boost score for sentences with keywords
            for keyword in keywords:
                if keyword.lower() in sentence.lower():
                    score += 10
            
            scored_sentences.append((score, sentence))
        
        # Sort by score and take top sentences
        scored_sentences.sort(key=lambda x: x[0], reverse=True)
        key_points = [sentence for _, sentence in scored_sentences[:max_points]]
        
        return key_points
    
    def get_json_depth(self, obj: Any, current_depth: int = 0) -> int:
        """Calculate the depth of a JSON object"""
        
        if isinstance(obj, dict):
            if not obj:
                return current_depth
            return max(self.get_json_depth(value, current_depth + 1) for value in obj.values())
        elif isinstance(obj, list):
            if not obj:
                return current_depth
            return max(self.get_json_depth(item, current_depth + 1) for item in obj)
        else:
            return current_depth

# Example usage and testing
if __name__ == "__main__":
    async def test_document_processor():
        processor = DocumentProcessor()
        
        # Test with sample text
        sample_text = "This is a test document. It contains multiple sentences. This helps test the processing capabilities."
        
        result = await processor.process_document(
            file_content=sample_text.encode('utf-8'),
            filename="test.txt",
            content_type="text/plain"
        )
        
        print(f"✅ Processed document: {result['filename']}")
        print(f"📄 Summary: {result['summary']}")
        print(f"🔑 Key points: {result['key_points']}")
    
    asyncio.run(test_document_processor())
